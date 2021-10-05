from operator import itemgetter

from docx import Document
from docx.shared import Cm, Pt, Inches, Mm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.shared import OxmlElement, qn

from ppb.models import FundRelease
import utils


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def item_sorter(items: list, *keys):
    return sorted(items, key=itemgetter(*keys))


def set_cell_margins(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')

    for m in [
        "top",
        "start",
        "bottom",
        "end",
    ]:
        if m in kwargs:
            node = OxmlElement("w:{}".format(m))
            node.set(qn('w:w'), str(kwargs.get(m)))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)

    tcPr.append(tcMar)


def indent_table(table, indent):
    tbl_pr = table._element.xpath('w:tblPr')
    if tbl_pr:
        e = OxmlElement('w:tblInd')
        e.set(qn('w:w'), str(indent))
        e.set(qn('w:type'), 'dxa')
        tbl_pr[0].append(e)

        e = OxmlElement('w:jc')
        e.set(qn('w:val'), str('center'))
        tbl_pr[0].append(e)


def create(rrf: FundRelease):
    document = Document()
    styles = document.styles
    font = styles['Normal'].font
    font.name = 'Arial'
    font.size = Pt(12)
    sections = document.sections
    header = sections[0].header
    footer = sections[0].footer

    paragraph = header.paragraphs[0]
    paragraph.text = 'By 2028, a world-class Army that is a source of national pride.'
    paragraph.paragraph_format.space_after = Cm(0)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph.runs[0].font.size = Pt(10)
    paragraph.runs[0].font.bold = True
    paragraph.runs[0].font.italic = True

    paragraph = footer.paragraphs[0]
    paragraph.text = '\nHonor. Patriotism. Duty'
    r = footer.add_paragraph(rrf.footer)
    r.runs[0].font.size = Pt(10)
    r.runs[0].font.italic = True

    r.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    r.paragraph_format.space_after = Cm(0)
    r.runs[0].font.size = Pt(10)

    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph.paragraph_format.space_after = Cm(0)
    paragraph.runs[0].font.size = Pt(10)
    paragraph.runs[0].font.bold = True
    paragraph.runs[0].font.italic = True

    r = footer.add_paragraph()
    r.paragraph_format.space_after = Cm(0)
    r.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    r = r.add_run()
    footer_pic = r.add_picture('ppb/static/ppb/footer.png')
    footer_pic.width = Mm(45)
    footer_pic.height = Mm(12)
    sections[0].footer_distance = Mm(5)
    sections[0].page_height = Mm(297)
    sections[0].page_width = Mm(210)

    for section in sections:
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Mm(7.5)

    page_table = document.add_table(rows=1, cols=1, style='Table Grid')
    page_table.rows[0].height = Cm(30)
    page = page_table.cell(0, 0)
    set_cell_margins(page, start=Cm(0), end=Cm(0))

    header_table = page.add_table(rows=1, cols=2)
    header_table.style = 'TableGrid'
    header_table.cell(0, 0).text = 'DISPOSITION FORM'
    r = header_table.cell(0, 0).paragraphs[0]
    r.paragraph_format.space_before = Pt(8)
    r.paragraph_format.space_after = Pt(8)

    header_table.cell(0, 1).text = 'SECURITY CLASSIFICATION (if any)'

    r.runs[0].font.size = Pt(20)
    r.runs[0].font.bold = True
    r.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    header_table = page.add_table(rows=1, cols=2)
    header_table.cell(0, 0).width = Inches(2.55)
    header_table.cell(0, 1).width = (header_table.cell(0, 1).width * 2 - header_table.cell(0, 0).width)
    header_table.style = 'TableGrid'
    header_table.cell(0, 0).text = 'FILE:'
    header_table.cell(0, 1).text = 'SUBJECT: '
    r = header_table.cell(0, 0).paragraphs[0].runs[0].font
    r.bold = True
    r = header_table.cell(0, 1).paragraphs[0]
    r = r.add_run('Request for the Release of Fund')
    r.bold = True

    delete_paragraph(page.paragraphs[0])
    p = page.paragraphs[0]
    p = p.add_run('.').font
    p.size = Pt(0)

    # page = page_table.cell(1, 0)  # NEW
    loc_paragraph = page.paragraphs[-1]
    loc_paragraph.text = 'TO: '
    p = loc_paragraph.add_run('C, OAFM')
    p.bold = True
    _ = loc_paragraph.add_run('\t\tFROM: ')
    p = loc_paragraph.add_run('G4')
    p.bold = True
    _ = loc_paragraph.add_run('\t\t\tDATE: ')
    _ = loc_paragraph.add_run('\t\tCMT Nr')
    cmt = f'\t\t\t\t\t\t\t\tMAJ BORROMEO/{rrf.created_by.account.initials.lower()}/6704\n'
    page.add_paragraph(cmt)

    item_one = f'\t1.\tReference:\tApproved SDF of OG4, PA dated____________: Subject: {rrf.reference}.'
    item_two = f'\t2.\tPer above reference, CGPA has approved the release of funds with following details:'
    item_three = f'\t3.\tAbove mentioned release shall be implemented in accordance with existing Government Accounting ' \
                 f'and Auditing Rules and Regulations and Government Procurement Law.'
    p = page.add_paragraph(item_one)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.paragraph_format.space_after = Pt(12)
    p = page.add_paragraph(item_two)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.paragraph_format.space_after = Pt(12)
    page.add_paragraph(f'\t\ta.\tMAJOR PAPs: {rrf.major_paps_print}')
    page.add_paragraph(f'\t\tb.\tSUGGESTED PAPs: {rrf.suggested_paps_print}')
    page.add_paragraph(f'\t\tc.\tSPECIFIC PAPs: {rrf.specific_paps_print}')
    p = page.add_paragraph(f'\t\td.\tAMOUNT: {utils.num2eng(rrf.amount).upper()} ')
    p = p.add_run(f'(PhP{rrf.amount:,.02f})')
    p.bold = True
    page.add_paragraph(f'\t\te.\tBREAKDOWN:\n')

    table = page.add_table(rows=1, cols=8)
    table.style = 'TableGrid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Servicing MFO'
    hdr_cells[1].text = 'Recipient Unit'
    hdr_cells[2].text = 'Program/ Project'
    hdr_cells[3].text = 'MA'
    hdr_cells[4].text = 'Object Code'
    hdr_cells[5].text = 'Amount (PhP)'
    hdr_cells[6].text = 'Specific Purpose'
    hdr_cells[7].text = 'Chargeability'

    hdr_cells[2].width = Mm(8)
    hdr_cells[3].width = Mm(20)
    hdr_cells[4].width = Mm(50)
    hdr_cells[5].width = Mm(38)
    hdr_cells[6].width = Mm(45)
    hdr_cells[7].width = Mm(30)

    old_table = None
    old_page = page

    for i in range(8):
        font = hdr_cells[i].paragraphs[0].runs[0].font
        font.size = Pt(8)
        font.bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    items = []
    for release_item in rrf.release_items.all():
        for recipient in release_item.release_recipients.all():
            item = {
                'servicing_mfo': recipient.servicing_mfo.name,
                'unit': recipient.unit.name,
                'program': release_item.program,
                'mission_area': release_item.mission_area.mission_area_group,
                'object_code': release_item.object_code.code,
                'amount': f'{recipient.amount:,.02f}',
                'specific_purpose': release_item.specific_purpose,
                'chargeability': release_item.chargeability
            }
            items.append(item)

    to_merge = [('servicing_mfo', 0), ('unit', 1), ('specific_purpose', 6), ('chargeability', 7)]
    merging = {}
    for item, row in to_merge:
        merging[item] = {'value': None, 'start': None, 'end': None, 'row': row}

    no_items = len(items)
    offset = False
    breaking_point = 1
    threshold = 570
    words = 0
    for i, item in enumerate(item_sorter(items, 'servicing_mfo', 'unit', 'program', 'mission_area'), start=1):
        row_cells = table.add_row().cells
        k = i

        if offset:
            k = i - breaking_point
            words = 0
        elif merging['specific_purpose']['start'] is None or \
                merging['specific_purpose']['start'] == merging['specific_purpose']['end']:
            words += len(item['specific_purpose'])
        else:
            words += 18

        if i == breaking_point and words < threshold:
            breaking_point += 1

        for key, pos in merging.items():
            if pos['start'] is None:
                pos['value'] = item[key]
                pos['start'] = table.cell(k, pos['row'])
                pos['end'] = table.cell(k, pos['row'])
                pos['start'].text = pos['value']
            elif pos['value'] == item[key]:
                pos['end'] = table.cell(k, pos['row'])
            else:
                if pos['start'] != pos['end']:
                    pos['start'].merge(pos['end'])

                pos['value'] = item[key]
                pos['start'] = table.cell(k, pos['row'])
                pos['end'] = table.cell(k, pos['row'])
                pos['start'].text = pos['value']

            if i == no_items and pos['start'] != pos['end']:
                pos['start'].merge(pos['end'])

            if not offset and i == breaking_point:
                if pos['start'] != pos['end']:
                    pos['start'].merge(pos['end'])
                pos['start'] = None

        row_cells[2].text = item['program']
        row_cells[3].text = item['mission_area']
        row_cells[4].text = item['object_code']
        row_cells[5].text = item['amount']

        for j in range(8):
            if len(row_cells[j].paragraphs[0].runs):
                font = row_cells[j].paragraphs[0].runs[0].font
                font.size = Pt(8)
            row_cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if j == 5:
                row_cells[j].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            elif j == 6 or j == 7:
                row_cells[j].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            else:
                row_cells[j].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        if not offset and (i >= breaking_point or words > threshold):
            offset = True
            page = document
            old_table = table
            table = page.add_table(rows=1, cols=8)
            table.style = 'TableGrid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Servicing MFO'
            hdr_cells[1].text = 'Recipient Unit'
            hdr_cells[2].text = 'Program/ Project'
            hdr_cells[3].text = 'MA'
            hdr_cells[4].text = 'Object Code'
            hdr_cells[5].text = 'Amount (PhP)'
            hdr_cells[6].text = 'Specific Purpose'
            hdr_cells[7].text = 'Chargeability'

            for r in range(8):
                font = hdr_cells[r].paragraphs[0].runs[0].font
                font.size = Pt(8)
                font.bold = True
                hdr_cells[r].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                hdr_cells[r].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if no_items > 1:
        total_cells = table.add_row().cells
        total_cells[0].merge(total_cells[4])
        total_cells[0].text = 'Total>>>>>'
        total_cells[5].text = f'{rrf.amount:,.02f}'

        for i in [0, 5]:
            total_font = total_cells[i].paragraphs[0].runs[0].font
            total_font.bold = True
            total_font.size = Pt(8)
            total_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            total_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if old_table is not None:
        indent_table(old_table, 36 * 4)
        hdr_cells = table.rows[0].cells
        hdr_cells[2].width = Mm(8)
        hdr_cells[3].width = Mm(20)
        hdr_cells[4].width = Mm(50)
        hdr_cells[5].width = Mm(38)
        hdr_cells[6].width = Mm(45)
        hdr_cells[7].width = Mm(30)
        space = '\n' * 3
    else:
        indent_table(table, 36 * 4)
        space = '\n' * 4

    if page != old_page:
        page.add_paragraph()

    threshold = 620

    if words + len(item_three) >= threshold:
        page = document
        words = 0
    else:
        words += len(item_three)

    p = page.add_paragraph(item_three)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.paragraph_format.space_after = Pt(12)

    item_four = '\t4.\tThis release shall be recorded as RRF No. '
    if words + len(item_four) + 27 >= threshold:
        page = document
        words = 0
    else:
        words += len(item_four) + 27
    p = page.add_paragraph(item_four)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    r = p.add_run(f'{rrf.rrf_no}')
    r.bold = True
    r.underline = True
    p.add_run(' for reference.')

    for p in old_page.paragraphs:
        p.paragraph_format.left_indent = Mm(2)
        p.paragraph_format.right_indent = Mm(2)

    signature_paragraph = page.add_paragraph(f'{space}ALEMANIA')
    signature_paragraph.runs[0].font.bold = True
    signature_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    return document
